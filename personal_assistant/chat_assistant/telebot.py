from dotenv import load_dotenv
import os
from aiogram import Bot, Dispatcher, executor, types
import sys
import time
from langchain_openai import ChatOpenAI  # Adjust path as per your project structure

from .info import groq_api_key, personal_prompt
from docx import Document

class Reference:
    '''
    This class is used to store the reference to the bot and the dispatcher
    '''
    
    def __init__(self) -> None:
        self.response=""


def clear_past_response():
    '''
    This function is used to clear the past response of the bot
    '''
    Reference.response=""

def save_to_docx(question, answer, response_time):
    '''
    Save question-answer pair and response time to a Word document.
    '''
    doc = Document()
    doc.add_heading('Question-Answer Pairs', 0)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Question'
    hdr_cells[1].text = 'Answer'
    hdr_cells[2].text = 'Response Time (seconds)'
    row = table.add_row().cells
    row[0].text = question
    row[1].text = answer
    row[2].text = str(response_time)
    doc.save('question_answer_pairs.docx')

@dispatcher.message_handler(commands=['start'])
async def welcome(message: types.Message):
    await message.reply("Welcome to this bot!\nPowered by aiogram.")

@dispatcher.message_handler(commands=['clear'])
async def welcome(message: types.Message):
    clear_past_response()
    await message.reply("The past response has been cleared.")

@dispatcher.message_handler(commands=['help'])
async def welcome(message: types.Message):
    await message.reply("Hey myself EVA , Press start so we can talk please !!.\n /start: Start the bot\n /help: Get help\n /clear: Clear the past response")



# Initialize the ChatOpenAI instance
llama3 = ChatOpenAI(
    api_key=groq_api_key,
    base_url="https://api.groq.com/openai/v1",
    model="llama3-8b-8192",
)



@dispatcher.message_handler()
async def all_time(message: types.Message):
    user_query=message.text
    print(f" >>> User: \t {query}")
    try:
        await message.answer("Please wait while I process your request...")
        start_time = time.time()

        full_query = f"User's name = {personal_prompt} +  {user_query}"
        query_msg = llama3.invoke(full_query)
        print("Before chatbot.chat call")
        bot_response=chatbot.chat(query)
        print("After chatbot.chat call")
        print(f"\n <<< Bot: \t\n {bot_response}")
        end_time = time.time()
       
        save_to_docx(user_query,bot_response,end_time-start_time)
    except Exception as e:  
        print(e)
        await message.answer("Sorry, I am not able to answer that.")

if __name__ == '__main__':
    executor.start_polling(Dispatcher, skip_updates=True)